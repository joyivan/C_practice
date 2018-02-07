
# coding: utf-8

# In[73]:


from docx import Document
from docx.shared import Inches
import docx
import os
print os.getcwd()
import re
import types


import sys
reload(sys)
sys.setdefaultencoding('utf8')
#sys.path.append("/Library/Frameworks/Python.framework/Versions/2.7/lib/python2.7/site-packages/docx")
def check_contain_chinese(check_str):
    for ch in check_str.decode('utf-8'):
        if u'\u4e00' <= ch <= u'\u9fff':
            print (ch)
    print ("END")
def split_str(sp_str):
    pattern = re.compile(r'\\t')
    print re.split(pattern,sp_str)

    
    
    
  



file=docx.Document("basicOfComputer.docx")
                
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

total=[]
#输出每一段的内容
qt=''
aa=''
ab=''
ac=''
ad=''
answers=[]
for para in file.paragraphs:
   # print(para.text)
    
    text=para.text
    #patternQ=re.compile(u'[^\u4E00-\u9FA5]')
    patternQ=re.compile(u'[\u4e00-\u9fa5]'|r'^ABCD')
    resultQ=re.match(patternQ,text)
    if resultQ:
        #print resultQ.group()+'TImu'
        
        #print '题目是:'+text
        qt=text
        
    
   # result=re.match
    patternA = re.compile(r'[ABCD]')
  
    resultA=re.match(patternA,text)
    
    if resultA:
        #print resultA.group()+'Daan'
        
        #timu.append(text.split())
        #patternEnd=re.compile(r'^C')
        #resultEnd=re.match(patternEnd,text)
        #if resultEnd:
        #    total.append(timu)
        #print str(timu).decode('utf-8')
    #print(str(timu).decode('unicode-escape'))
        
        if text.startswith('A、') or text.startswith('C、'):
            answers=text.split()
            if len(answers)==2:
                if text.startswith('A、'):
                    
                    aa=answers[0]
                    ab=answers[1]
                if text.startswith('C、'):
                    ac=answers[0]
                    ad=answers[1]
                    total.append([qt,aa,ab,ac,ad])
                    qt=''
                    aa=''
                    ab=''
                    ac=''
                    ad=''
            if len(answers)==4:
                aa=answers[0]
                ab=answers[1]
                ac=answers[2]
                ad=answers[3]
            
                total.append([qt,aa,ab,ac,ad])
                qt=''
                aa=''
                ab=''
                ac=''
                ad=''
           
            
            
    #patternC=re.compile(r'^C')
    #resultC=re.match(patternC,text)
    #if resultC:
    #    total.append(timu)
    #    timu=[]
    print 'close'
    
   
        
   
   # check_contain_chinese(para.text) 

#输出段落编号及段落内容
#for i in range(len(file.paragraphs)):
#    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)


# In[68]:


import xlwt
from datetime import datetime
os.remove('example.xls')
style0 = xlwt.easyxf('font: name Times New Roman, color-index red, bold on',
    num_format_str='#,##0.00')
style1 = xlwt.easyxf(num_format_str='D-MMM-YY')

wb = xlwt.Workbook()
ws = wb.add_sheet('A Test Sheet')

for i in range(len(total)):
    for j in range(len(total[i])):
        ws.write(i, j,total[i][j], style0)
        

wb.save('example.xls')



# In[44]:

print(str(total).decode('unicode-escape'))
print


# In[36]:


import numpy as np
for i in range(10):
    print i
    
    


# In[45]:

len(total)


# In[48]:

total


# In[60]:

os.remove('example.xls')


# In[ ]:



